"""Tailors resumes per job using Azure OpenAI."""

import json
import logging
import os
from datetime import datetime

from docx import Document
from openai import AsyncAzureOpenAI
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import settings
from app.models.models import JobListing, JobStatus, TailoredResume, UserProfile

logger = logging.getLogger(__name__)

RESUME_DIR = os.path.join(os.path.dirname(__file__), "..", "..", "resumes")

SYSTEM_PROMPT = """You are an expert resume writer. Given a candidate's base resume and a job description,
rewrite the resume to better match the job. Guidelines:
- Keep all information factual — do NOT fabricate experience, skills, or credentials
- Reorder and emphasize bullet points that are most relevant to the JD
- Mirror keywords and phrases from the JD naturally
- Keep the resume to approximately the same length
- Return the full resume text (plain text, not markdown)
- After the resume, add a section labeled "COVER_LETTER:" with a brief, tailored cover letter (3-4 paragraphs)"""


async def tailor_resume_for_job(
    client: AsyncAzureOpenAI, base_resume_text: str, job: JobListing
) -> dict:
    """Generate a tailored resume and cover letter for a job."""
    prompt = f"""## Base Resume
{base_resume_text}

## Target Job
- Title: {job.title}
- Company: (from job listing)
- Location: {job.location}
- Department: {job.department}

## Job Description
{job.description_text[:4000]}

Rewrite the resume to optimize for this specific role. Then provide a cover letter."""

    response = await client.chat.completions.create(
        model=settings.azure_openai_deployment,
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": prompt},
        ],
        temperature=0.4,
    )
    content = response.choices[0].message.content

    # Split resume and cover letter
    if "COVER_LETTER:" in content:
        parts = content.split("COVER_LETTER:", 1)
        resume_text = parts[0].strip()
        cover_letter = parts[1].strip()
    else:
        resume_text = content.strip()
        cover_letter = ""

    return {"resume_text": resume_text, "cover_letter": cover_letter}


def _save_resume_as_docx(resume_text: str, filename: str) -> str:
    """Save resume text as a DOCX file and return the path."""
    os.makedirs(RESUME_DIR, exist_ok=True)
    filepath = os.path.join(RESUME_DIR, filename)
    doc = Document()
    for line in resume_text.split("\n"):
        doc.add_paragraph(line)
    doc.save(filepath)
    return filepath


def _read_base_resume(path: str) -> str:
    """Read a base resume from a DOCX or TXT file."""
    if path.endswith(".docx"):
        doc = Document(path)
        return "\n".join(p.text for p in doc.paragraphs)
    else:
        with open(path, "r", encoding="utf-8") as f:
            return f.read()


async def tailor_resumes_for_matched_jobs(db: AsyncSession) -> int:
    """Tailor resumes for all matched jobs that don't have one yet. Returns count."""
    # Get user profile
    result = await db.execute(select(UserProfile).limit(1))
    profile = result.scalar_one_or_none()
    if not profile or not profile.base_resume_path:
        logger.warning("No user profile or base resume found — skipping tailoring")
        return 0

    base_resume = _read_base_resume(profile.base_resume_path)

    # Get matched jobs without tailored resumes
    result = await db.execute(
        select(JobListing)
        .where(JobListing.status == JobStatus.MATCHED)
        .outerjoin(TailoredResume)
        .where(TailoredResume.id.is_(None))
    )
    jobs = result.scalars().all()
    if not jobs:
        logger.info("No matched jobs needing resume tailoring")
        return 0

    client = AsyncAzureOpenAI(
        api_key=settings.azure_openai_api_key,
        azure_endpoint=settings.azure_openai_endpoint,
        api_version=settings.azure_openai_api_version,
    )

    tailored = 0
    for job in jobs:
        try:
            result_data = await tailor_resume_for_job(client, base_resume, job)
            filename = f"resume_job_{job.id}_{datetime.utcnow().strftime('%Y%m%d%H%M%S')}.docx"
            pdf_path = _save_resume_as_docx(result_data["resume_text"], filename)

            tr = TailoredResume(
                job_id=job.id,
                resume_content=result_data["resume_text"],
                resume_pdf_path=pdf_path,
                cover_letter=result_data["cover_letter"],
            )
            db.add(tr)
            tailored += 1
        except Exception as e:
            logger.error(f"Error tailoring resume for job {job.id}: {e}")

    await db.commit()
    logger.info(f"Tailored {tailored} resumes")
    return tailored
